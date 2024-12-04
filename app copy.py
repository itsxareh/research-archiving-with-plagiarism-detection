import os
import re
from typing import Dict, List, Tuple, Any
from flask import Flask, request, jsonify, send_from_directory
from PyPDF2 import PdfReader
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import mysql.connector
from mysql.connector import Error
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from datetime import datetime
from threading import Thread
import pytz
import numpy as np
import string
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')
nltk.download('averaged_perceptron_tagger')

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024

def get_db():
    try:
        conn = mysql.connector.connect(
            host='localhost',
            user='root', 
            password='', 
            database='research_repository' 
        )
        if conn.is_connected():
            print ("Connected to MySQL")
            return conn
    except Error as e:
        print(f"Error connecting to MySQL: {e}")
        return None
    
def init_db():
    with get_db() as conn:
        cursor = conn.cursor()
        
        # Modified archive_research table
        cursor.execute('''CREATE TABLE IF NOT EXISTS archive_research (
            id INT AUTO_INCREMENT PRIMARY KEY,
            archive_id INT NOT NULL,
            student_id VARCHAR(20) NOT NULL,
            department_id INT NOT NULL,
            course_id INT NOT NULL,
            project_title VARCHAR(255) NOT NULL,
            dateOfSubmit DATETIME NOT NULL,
            project_year VARCHAR(4) NOT NULL,
            project_abstract TEXT,
            keywords TEXT,
            content TEXT NOT NULL,
            research_owner_email VARCHAR(255),
            project_members TEXT,
            documents VARCHAR(255),
            file_size BIGINT UNSIGNED,
            page_count INT UNSIGNED,
            word_count INT UNSIGNED,
            character_count INT UNSIGNED
        )''')

        cursor.execute('''CREATE TABLE IF NOT EXISTS plagiarism_results (
            id INT AUTO_INCREMENT PRIMARY KEY,
            archive_id INT NOT NULL,
            submitted_sentence TEXT NOT NULL,
            existing_sentence TEXT NOT NULL,
            similar_archive_id INT NOT NULL,
            similarity_percentage DECIMAL(5,2) NOT NULL,
            is_plagiarized BOOLEAN NOT NULL,
            FOREIGN KEY (archive_id) REFERENCES archive_research(id),
            FOREIGN KEY (similar_archive_id) REFERENCES archive_research(id)
        )''')

        cursor.execute('''CREATE TABLE IF NOT EXISTS plagiarism_summary (
            id INT AUTO_INCREMENT PRIMARY KEY,
            archive_id INT NOT NULL,
            similar_archive_id INT NOT NULL,
            plagiarism_percentage DECIMAL(5,2) NOT NULL,
            FOREIGN KEY (archive_id) REFERENCES archive_research(id),
            FOREIGN KEY (similar_archive_id) REFERENCES archive_research(id)
        )''')
        
        conn.commit()

def get_file_metrics(file_path, file_type):
    """Calculate file metrics including size, page count, word count, and character count"""
    metrics = {
        'file_size': os.path.getsize(file_path),
        'page_count': 0,
        'word_count': 0,
        'character_count': 0
    }
    
    content = ""
    if file_type == 'application/pdf':
        reader = PdfReader(file_path)
        metrics['page_count'] = len(reader.pages)
        for page in reader.pages:
            page_text = page.extract_text()
            content += page_text
            
    elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        doc = Document(file_path)
        metrics['page_count'] = len(doc.paragraphs) // 40  # Approximate pages
        for para in doc.paragraphs:
            content += para.text + "\n"
            
    else:  # For text files
        with open(file_path, 'r') as file:
            content = file.read()
            metrics['page_count'] = len(content) // 3000  # Approximate pages
    
    # Calculate word and character counts
    metrics['word_count'] = len(content.split())
    metrics['character_count'] = len(content)
    
    return metrics, content

def normalize_text(text):
    # Convert to lowercase
    text = text.lower()

    # Remove leading and trailing spaces
    text = text.strip()
    
    # Replace unicode punctuation with standard punctuation
    text = text.translate(str.maketrans('', '', string.punctuation))
    
    # Remove multiple spaces between words
    text = re.sub(r'\s+', ' ', text)
    
    # Normalize hyphenated words (removes spaces around hyphens)    
    text = re.sub(r'\s*-\s*', '-', text)
    
    # Remove non-alphanumeric characters except spaces and hyphens
    text = re.sub(r'[^\w\s-]', '', text)

    # Handle redundant punctuation like multiple dots
    text = re.sub(r'\.\.+', '.', text)

    # Replace multiple consecutive hyphens with a single hyphen
    text = re.sub(r'-{2,}', '-', text)
    
    return text

def split_into_sentences(text):
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt')
    
    sentences =  sent_tokenize(text)

    filtered_sentences = [
        sentence for sentence in sentences
        if 8 <= len(sentence.split()) <= 50
    ]
    return filtered_sentences

def calculate_similarity(submitted_content, submissions, similarity_threshold=0.7):
    submitted_sentences = [normalize_text(sentence) for sentence in split_into_sentences(submitted_content)]

    # Prepare a list of sentences and their corresponding submission IDs
    all_existing_sentences = []
    all_existing_submission_ids = []

    for submission_id, content in submissions:
        sentences = [normalize_text(sentence) for sentence in split_into_sentences(content)]
        all_existing_sentences.extend(sentences)
        all_existing_submission_ids.extend([submission_id] * len(sentences))  # Match each sentence with its submission ID

    # Combine submitted and existing sentences
    all_sentences = submitted_sentences + all_existing_sentences

    if not all_sentences:
        return []

    # Vectorize the sentences
    vectorizer = TfidfVectorizer(lowercase=True, stop_words=None, token_pattern=r'\b\w+\b').fit_transform(all_sentences)

    # Similarity between submitted sentences and all existing ones
    similarity_matrix = cosine_similarity(vectorizer[:len(submitted_sentences)], vectorizer[len(submitted_sentences):])

    # Find matching sentences
    matching_results = []
    for i, submitted_sentence in enumerate(submitted_sentences):
        for j, existing_sentence in enumerate(all_existing_sentences):
            similarity_score = similarity_matrix[i][j]
            if similarity_score > similarity_threshold:
                matching_results.append({
                    'submitted_sentence': submitted_sentence,
                    'existing_sentence': existing_sentence,
                    'similarity_percentage': similarity_score * 100,
                    'similar_archive_id': all_existing_submission_ids[j]
                })

    return matching_results

@app.route('/')
def index():    
    return send_from_directory('.', 'student/project_list.php')

@app.route('/upload_research', methods=['POST'])
def upload_file():
    print("Form Data:", request.form)
    print("Files Data:", request.files)
    if 'file' not in request.files:
        return jsonify({'status': 'error', 'message': "No file part"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'status': 'error', 'message': "No selected file"}), 400

    # Extract form data
    archive_id = request.form.get('archive_id')
    student_id = request.form.get('student_id')
    project_title = request.form.get('project_title')
    project_year = request.form.get('year')
    department_id = request.form.get('department_id')
    course_id = request.form.get('course_id')
    abstract = request.form.get('abstract')
    keywords = request.form.get('keywords')
    project_members = request.form.get('project_members')
    pdf_path = request.form.get('pdf_path')
    owner_email = request.form.get('owner_email')

    if not student_id:
        return jsonify({'status': 'error', 'message': "No student ID provided"}), 400

    try:
        # Step 1: Save the file
        upload_folder = os.path.join(os.getcwd(), 'pdf_files')
        os.makedirs(upload_folder, exist_ok=True)
        file_path = os.path.join(upload_folder, file.filename)

        # File metrics and content
        file_metrics, content = get_file_metrics(file_path, file.content_type)
        if not content:
            return jsonify({'status': 'error', 'message': "Failed to extract content from file"}), 400

        # Save the uploaded file details in the database
        manila_tz = pytz.timezone('Asia/Manila')
        current_time = datetime.now(manila_tz)
        
        with get_db() as conn:
            cursor = conn.cursor()
            # First, insert the research with initial status
            cursor.execute("""
                INSERT INTO archive_research (
                    archive_id, student_id, department_id, course_id,
                    project_title, dateOfSubmit, project_year, project_abstract,
                    keywords, content, research_owner_email, project_members,
                    documents, file_size, page_count, word_count, character_count,
                    document_status  -- Added initial status
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    archive_id, student_id, department_id, course_id,
                    project_title, current_time, project_year, abstract,
                    keywords, content, owner_email, project_members,
                    pdf_path, file_metrics['file_size'], file_metrics['page_count'],
                    file_metrics['word_count'], file_metrics['character_count'],
                    'Processing'  # Initial status while plagiarism check runs
                ))
            conn.commit()
            new_archive_id = cursor.lastrowid

            # Get department info for response
            cursor.execute("SELECT * FROM departments WHERE id = %s", (department_id,))
            departments_db = cursor.fetchall()
            department = departments_db[0][2]

        # Prepare response before starting plagiarism check
        response = {
            'status': 'success',
            'department': department,
            'document_status': 'Processing',  # Initial status
            'file_metrics': file_metrics,
            'archive_id': archive_id,
            'message': 'File uploaded successfully. Plagiarism check is running in the background.'
        }

        plagiarism_check(new_archive_id, student_id, content, current_time)

        return jsonify(response)

    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

def plagiarism_check(new_archive_id, student_id, content, current_time):
    """Separate function for background plagiarism checking"""
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            
            # Update status to show plagiarism check is in progress
            cursor.execute("""
                UPDATE archive_research 
                SET document_status = 'Checking Plagiarism' 
                WHERE id = %s
            """, (new_archive_id,))
            conn.commit()

            # Fetch existing submissions for comparison
            cursor.execute("SELECT id, content FROM archive_research WHERE student_id != %s", (student_id,))
            submissions = cursor.fetchall()

            plagiarism_results = {}
            if submissions:
                matching_sentences = calculate_similarity(content, submissions)
                for match in matching_sentences:
                    similar_id = match['similar_archive_id']

                    cursor.execute("SELECT content FROM archive_research WHERE id = %s", (similar_id,))
                    similar_content = cursor.fetchone()[0]
                    total_sentences = len(split_into_sentences(similar_content))

                    cursor.execute("""
                        INSERT INTO plagiarism_results (
                            archive_id, similar_archive_id, submitted_sentence,
                            existing_sentence, similarity_percentage, is_plagiarized
                        ) VALUES (%s, %s, %s, %s, %s, %s)
                        """, (
                        new_archive_id, similar_id, match['submitted_sentence'],
                        match['existing_sentence'], match['similarity_percentage'], True
                    ))

                    if similar_id not in plagiarism_results:
                        plagiarism_results[similar_id] = {
                            'matched_sentences': 0,
                            'total_similarity': 0,
                            'total_sentences': total_sentences
                        }
                    plagiarism_results[similar_id]['matched_sentences'] += 1
                    plagiarism_results[similar_id]['total_similarity'] += match['similarity_percentage']

                for similar_id, results in plagiarism_results.items():
                    avg_similarity = results['total_similarity'] / results['matched_sentences']
                    plagiarism_percentage = (results['total_similarity'] / results['total_sentences']) if results['total_sentences'] > 0 else 0

                    cursor.execute("""
                        INSERT INTO plagiarism_summary (
                            archive_id, similar_archive_id, plagiarism_percentage
                        ) VALUES (%s, %s, %s)
                        """, (new_archive_id, similar_id, plagiarism_percentage))
            conn.commit()

            print(plagiarism_results)
            # Calculate final plagiarism results and update status
            cursor.execute("""
                SELECT COUNT(plagiarism_summary.id) as total_ids, 
                       SUM(plagiarism_percentage) as total_percentage 
                FROM plagiarism_summary 
                WHERE archive_id = %s
                GROUP BY archive_id
                """, (new_archive_id,))
            
            result = cursor.fetchone()
            sum_of_percentage = result[1] if result else 0
            document_status = "Accepted" if sum_of_percentage < 20 else "Not Accepted"
            date_publish = current_time.strftime('%Y-%m-%d') if document_status == "Accepted" else None

            cursor.execute("""
                UPDATE archive_research 
                SET document_status = %s, 
                    date_published = %s 
                WHERE id = %s
                """, (document_status, date_publish, new_archive_id))
            conn.commit()

    except Exception as e:
        print(f"Error during plagiarism check: {str(e)}")

@app.route('/check_db_connection', methods=['GET'])
def check_db_connection():
    try:
        with get_db() as conn:
            conn.execute("SELECT 1")
        return jsonify(status="Database connection successful"), 200
    except Exception as e:
        return jsonify(status="Database connection failed", error=str(e)), 500
    
if __name__ == '__main__':
    init_db()
    app.run(debug=True, port=3000)