import os
import re
from typing import Dict, List, Tuple, Any
from flask import Flask, request, jsonify, send_from_directory
from PyPDF2 import PdfReader
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import mysql.connector
from mysql.connector import Error
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from datetime import datetime
import pytz
import numpy as np

nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')
nltk.download('averaged_perceptron_tagger')

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024

def get_db():
    try:
        conn = mysql.connector.connect(
            host='localhost',
            user='root', 
            password='', 
            database='online_thesis' 
        )
        if conn.is_connected():
            print ("Connected to MySQL")
            return conn
    except Error as e:
        print(f"Error connecting to MySQL: {e}")
        return None
    
def init_db():
    with get_db() as conn:
        cursor = conn.cursor()
        
        # Modified archive_research table
        cursor.execute('''CREATE TABLE IF NOT EXISTS archive_research (
            id INT AUTO_INCREMENT PRIMARY KEY,
            archive_id INT NOT NULL,
            student_id VARCHAR(20) NOT NULL,
            department_id INT NOT NULL,
            course_id INT NOT NULL,
            project_title VARCHAR(255) NOT NULL,
            dateOfSubmit DATETIME NOT NULL,
            project_year VARCHAR(4) NOT NULL,
            project_abstract TEXT,
            keywords TEXT,
            content TEXT NOT NULL,
            research_owner_email VARCHAR(255),
            project_members TEXT,
            documents VARCHAR(255),
            file_size BIGINT UNSIGNED,
            page_count INT UNSIGNED,
            word_count INT UNSIGNED,
            character_count INT UNSIGNED
        )''')

        cursor.execute('''CREATE TABLE IF NOT EXISTS plagiarism_results (
            id INT AUTO_INCREMENT PRIMARY KEY,
            archive_id INT NOT NULL,
            submitted_sentence TEXT NOT NULL,
            existing_sentence TEXT NOT NULL,
            similar_archive_id INT NOT NULL,
            similarity_percentage DECIMAL(5,2) NOT NULL,
            is_plagiarized BOOLEAN NOT NULL,
            FOREIGN KEY (archive_id) REFERENCES archive_research(id),
            FOREIGN KEY (similar_archive_id) REFERENCES archive_research(id)
        )''')

        cursor.execute('''CREATE TABLE IF NOT EXISTS plagiarism_summary (
            id INT AUTO_INCREMENT PRIMARY KEY,
            archive_id INT NOT NULL,
            similar_archive_id INT NOT NULL,
            plagiarism_percentage DECIMAL(5,2) NOT NULL,
            FOREIGN KEY (archive_id) REFERENCES archive_research(id),
            FOREIGN KEY (similar_archive_id) REFERENCES archive_research(id)
        )''')
        
        conn.commit()

class TextPreprocessor:
    def __init__(self):
        self.lemmatizer = WordNetLemmatizer()
        self.stop_words = set(stopwords.words('english'))
        
    def clean_text(self, text: str) -> str:
        """Remove special characters and normalize text"""
        # Convert to lowercase
        text = text.lower()
        
        # Remove citations and references
        text = re.sub(r'\[\d+\]', '', text)
        text = re.sub(r'\(\w+,\s*\d{4}\)', '', text)
        
        # Remove URLs
        text = re.sub(r'http\S+|www\S+|https\S+', '', text)
        
        # Remove email addresses
        text = re.sub(r'\S+@\S+', '', text)
        
        # Remove special characters but keep sentence structure
        text = re.sub(r'[^\w\s\.\?\!]', '', text)
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    def lemmatize_text(self, text: str) -> str:
        """Lemmatize text while preserving sentence structure"""
        words = word_tokenize(text)
        lemmatized_words = [self.lemmatizer.lemmatize(word) for word in words if word not in self.stop_words]
        return ' '.join(lemmatized_words)

class DocumentMetrics:
    @staticmethod
    def get_file_metrics(file_path: str, file_type: str) -> Tuple[Dict[str, int], str]:
        """Calculate enhanced file metrics and extract content"""
        metrics = {
            'file_size': os.path.getsize(file_path),
            'page_count': 0,
            'word_count': 0,
            'character_count': 0,
            'sentence_count': 0,
            'paragraph_count': 0
        }
        print(file_path)
        print(metrics)
        content = ""
        
        try:
            if file_type == 'application/pdf':
                content, metrics = DocumentMetrics._process_pdf(file_path)
            elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                content, metrics = DocumentMetrics._process_docx(file_path)
            else:
                content, metrics = DocumentMetrics._process_text(file_path)
                
            # Calculate additional metrics
            metrics['sentence_count'] = len(sent_tokenize(content))
            metrics['word_count'] = len(word_tokenize(content))
            metrics['character_count'] = len(content)
            
        except Exception as e:
            print(f"Error processing file: {str(e)}")
            
        return metrics, content
    
    @staticmethod
    def _process_pdf(file_path: str) -> Tuple[str, Dict[str, int]]:
        reader = PdfReader(file_path)
        content = ""
        metrics = {'page_count': len(reader.pages), 'paragraph_count': 0, 'file_size': os.path.getsize(file_path)}
        
        for page in reader.pages:
            page_text = page.extract_text()
            paragraphs = page_text.split('\n\n')
            metrics['paragraph_count'] += len(paragraphs)
            content += page_text
            
        return content, metrics
    
    @staticmethod
    def _process_docx(file_path: str) -> Tuple[str, Dict[str, int]]:
        doc = Document(file_path)
        content = ""
        metrics = {'page_count': 0, 'paragraph_count': len(doc.paragraphs)}
        
        for para in doc.paragraphs:
            content += para.text + "\n"
        
        # Estimate page count based on standard page length
        metrics['page_count'] = max(1, len(content) // 3000)
        
        return content, metrics
    
    @staticmethod
    def _process_text(file_path: str) -> Tuple[str, Dict[str, int]]:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            
        paragraphs = content.split('\n\n')
        metrics = {
            'page_count': max(1, len(content) // 3000),
            'paragraph_count': len(paragraphs)
        }
        
        return content, metrics

class PlagiarismDetector:
    def __init__(self, threshold: float = 0.8):
        self.preprocessor = TextPreprocessor()
        self.similarity_threshold = threshold
        self.vectorizer = TfidfVectorizer(
            lowercase=True,
            token_pattern=r'\b\w+\b',
            ngram_range=(1, 3),  # Consider up to trigrams
            max_features=10000
        )
        
    def split_into_chunks(self, text: str) -> List[Tuple[str, List[str]]]:
        """Split text into meaningful chunks and keep track of original sentences"""
        sentences = sent_tokenize(text)
        chunks = []
        chunk_sentences = []  # Keep track of which sentences belong to which chunk
        current_chunk = []
        current_chunk_sentences = []
        current_length = 0
        
        for sentence in sentences:
            words = word_tokenize(sentence)
            if current_length + len(words) > 50:  # Aim for ~50 words per chunk
                if current_chunk:
                    chunks.append(' '.join(current_chunk))
                    chunk_sentences.append(current_chunk_sentences)
                current_chunk = [sentence]
                current_chunk_sentences = [sentence]
                current_length = len(words)
            else:
                current_chunk.append(sentence)
                current_chunk_sentences.append(sentence)
                current_length += len(words)
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
            chunk_sentences.append(current_chunk_sentences)
            
        # Filter out chunks that are too short
        valid_chunks = []
        valid_sentences = []
        for chunk, sents in zip(chunks, chunk_sentences):
            if len(word_tokenize(chunk)) >= 7:
                valid_chunks.append(chunk)
                valid_sentences.append(sents)
                
        return list(zip(valid_chunks, valid_sentences))
    
    def calculate_similarity(self, submitted_content: str, submissions: List[Tuple[int, str]]) -> List[Dict[str, Any]]:
        """Calculate similarity with enhanced detection and proper sentence tracking"""
        # Preprocess submitted content
        cleaned_submitted = self.preprocessor.clean_text(submitted_content)
        lemmatized_submitted = self.preprocessor.lemmatize_text(cleaned_submitted)
        submitted_chunks = self.split_into_chunks(lemmatized_submitted)
        
        # Prepare existing submissions
        all_existing_chunks = []
        all_existing_sentences = []
        all_existing_submission_ids = []
        
        for submission_id, content in submissions:
            cleaned_content = self.preprocessor.clean_text(content)
            lemmatized_content = self.preprocessor.lemmatize_text(cleaned_content)
            chunks_and_sentences = self.split_into_chunks(lemmatized_content)
            
            for chunk, sentences in chunks_and_sentences:
                all_existing_chunks.append(chunk)
                all_existing_sentences.append(sentences)
                all_existing_submission_ids.append(submission_id)
        
        if not all_existing_chunks or not submitted_chunks:
            return []
        
        # Vectorize all chunks
        all_chunks = [chunk for chunk, _ in submitted_chunks] + all_existing_chunks
        try:
            vectors = self.vectorizer.fit_transform(all_chunks)
            similarity_matrix = cosine_similarity(
                vectors[:len(submitted_chunks)],
                vectors[len(submitted_chunks):]
            )
        except Exception as e:
            print(f"Vectorization error: {str(e)}")
            return []
        
        # Find matching chunks with proper sentence tracking
        matching_results = []
        for i, (submitted_chunk, submitted_sentences) in enumerate(submitted_chunks):
            for j, existing_sentences in enumerate(all_existing_sentences):
                similarity_score = similarity_matrix[i][j]
                
                if similarity_score > self.similarity_threshold:
                    # Use the first sentence from each chunk for display
                    matching_results.append({
                        'submitted_chunk': submitted_sentences[0],
                        'existing_chunk': existing_sentences[0],
                        'similarity_percentage': similarity_score * 100,
                        'similar_archive_id': all_existing_submission_ids[j],
                        'word_count': len(word_tokenize(submitted_sentences[0]))
                    })
        
        return matching_results

# Helper function to use the enhanced system
def process_document_for_plagiarism(file_path: str, file_type: str, submissions: List[Tuple[int, str]], similarity_threshold: float = 0.8) -> Tuple[Dict[str, int], List[Dict[str, Any]]]:
    """Process a document and check for plagiarism with enhanced detection"""
    # Get document metrics
    metrics, content = DocumentMetrics.get_file_metrics(file_path, file_type)
    
    # Initialize plagiarism detector
    detector = PlagiarismDetector(threshold=similarity_threshold)
    
    # Get plagiarism results
    plagiarism_results = detector.calculate_similarity(content, submissions)
    
    return metrics, content, plagiarism_results

@app.route('/')
def index():    
    return send_from_directory('.', 'student/project_list.php')

@app.route('/upload_research', methods=['POST'])
def upload_file():
    try:
        # Basic request validation
        if 'file' not in request.files:
            return jsonify(error="No file part"), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify(error="No selected file"), 400

        # Extract form data
        form_data = {
            'archive_id': request.form.get('archive_id'),
            'student_id': request.form.get('student_id'),
            'project_title': request.form.get('project_title'),
            'date_of_submit': request.form.get('date_of_submit'),
            'project_year': request.form.get('year'),
            'department_id': request.form.get('department_id'),
            'course_id': request.form.get('course_id'),
            'abstract': request.form.get('abstract'),
            'keywords': request.form.get('keywords'),
            'project_members': request.form.get('project_members'),
            'pdf_path': request.form.get('pdf_path'),
            'owner_email': request.form.get('owner_email')
        }

        if not form_data['student_id']:
            return jsonify(error="No student ID provided"), 400

        # Save uploaded file
        upload_folder = os.path.join(os.getcwd(), 'pdf_files')
        os.makedirs(upload_folder, exist_ok=True)
        file_path = os.path.join(upload_folder, file.filename)
        file.save(file_path)

        # Get existing submissions for comparison
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT id, content FROM archive_research WHERE student_id != %s",
                (form_data['student_id'],)
            )
            existing_submissions = cursor.fetchall()

        # Process document and check for plagiarism
        file_metrics, content, plagiarism_matches = process_document_for_plagiarism(
            file_path=file_path,
            file_type=file.content_type,
            submissions=existing_submissions,
            similarity_threshold=0.7  # Can be adjusted based on requirements
        )

        # Store results in database
        with get_db() as conn:
            cursor = conn.cursor()
            
            # Insert into archive_research
            manila_tz = pytz.timezone('Asia/Manila')
            current_time = datetime.now(manila_tz)
            cursor.execute("""
                INSERT INTO archive_research (
                    archive_id, student_id, department_id, course_id,
                    project_title, dateOfSubmit, project_year, project_abstract,
                    keywords, content, research_owner_email, project_members,
                    documents, file_size, page_count, word_count, character_count
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    form_data['archive_id'], form_data['student_id'],
                    form_data['department_id'], form_data['course_id'],
                    form_data['project_title'], current_time,
                    form_data['project_year'], form_data['abstract'],
                    form_data['keywords'], content,
                    form_data['owner_email'], form_data['project_members'],
                    form_data['pdf_path'], file_metrics['file_size'],
                    file_metrics['page_count'], file_metrics['word_count'],
                    file_metrics['character_count']
                )
            )
            conn.commit()
            new_archive_id = cursor.lastrowid

            # Process and store plagiarism results
            plagiarism_summary = {}
            for match in plagiarism_matches:
                similar_id = match['similar_archive_id']
                
                # Store individual matches
                cursor.execute("""
                    INSERT INTO plagiarism_results (
                        archive_id, similar_archive_id, submitted_sentence,
                        existing_sentence, similarity_percentage, is_plagiarized
                    ) VALUES (%s, %s, %s, %s, %s, %s)
                    """, (
                        new_archive_id, similar_id,
                        match['submitted_chunk'],
                        match['existing_chunk'],
                        match['similarity_percentage'],
                        True
                    )
                )

                # Aggregate results for summary
                if similar_id not in plagiarism_summary:
                    plagiarism_summary[similar_id] = {
                        'matched_chunks': 0,
                        'total_similarity': 0,
                        'total_words': 0
                    }
                plagiarism_summary[similar_id]['matched_chunks'] += 1
                plagiarism_summary[similar_id]['total_similarity'] += match['similarity_percentage']
                plagiarism_summary[similar_id]['total_words'] += match['word_count']

            # Store plagiarism summary
            for similar_id, summary in plagiarism_summary.items():
                avg_similarity = summary['total_similarity'] / summary['matched_chunks']
                cursor.execute("""
                    INSERT INTO plagiarism_summary (
                        archive_id, similar_archive_id, plagiarism_percentage
                    ) VALUES (%s, %s, %s)
                    """, (
                        new_archive_id, similar_id, avg_similarity
                    )
                )

            conn.commit()

        # Prepare response
        response_data = {
            'status': 'success',
            'file_metrics': file_metrics,
            'plagiarism_results': [
                {
                    'similar_archive_id': similar_id,
                    'plagiarism_percentage': summary['total_similarity'] / summary['matched_chunks'],
                    'matched_chunks': summary['matched_chunks'],
                    'total_words': summary['total_words']
                }
                for similar_id, summary in plagiarism_summary.items()
            ]
        }

        return jsonify(response_data)

    except Exception as e:
        print(f"Error processing upload: {str(e)}")
        return jsonify(error=str(e)), 500
    
@app.route('/check_db_connection', methods=['GET'])
def check_db_connection():
    try:
        with get_db() as conn:
            conn.execute("SELECT 1")
        return jsonify(status="Database connection successful"), 200
    except Exception as e:
        return jsonify(status="Database connection failed", error=str(e)), 500
    
if __name__ == '__main__':
    init_db()
    app.run(debug=True, port=3000)