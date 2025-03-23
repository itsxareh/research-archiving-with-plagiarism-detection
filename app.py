import os
import re
from typing import Dict, List, Tuple, Any
from flask import Flask, request, jsonify, send_from_directory
from PyPDF2 import PdfReader
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import mysql.connector
from mysql.connector import Error
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from datetime import datetime
from threading import Thread
import pytesseract
#https://github.com/oschwartz10612/poppler-windows/releases/
#https://github.com/UB-Mannheim/tesseract/wiki
#https://github.com/tesseract-ocr/tessdata
if os.name == 'nt':  # for Windows
    POPPLER_PATH = r"C:\Program Files\poppler-24.08.0\Library\bin"
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
try:
    print(f"Tesseract Version: {pytesseract.get_tesseract_version()}")
    print(f"Available languages: {pytesseract.get_languages()}")
except Exception as e:
    print(f"Tesseract Error: {str(e)}")
    print("Please ensure Tesseract is properly installed and the path is correct.")

from PIL import Image, ImageEnhance
from pdf2image import convert_from_path
import cv2
import io
import traceback
import fitz 
import pytz
import numpy as np
import string
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')
nltk.download('averaged_perceptron_tagger')

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024

def get_db():
    try:
        conn = mysql.connector.connect(
            host='localhost',
            user='root', 
            password='', 
            database='research_repository' 
        )
        if conn.is_connected():
            return conn
    except Error as e:
        print(f"Error connecting to MySQL: {e}")
        return None
    
def init_db():
    with get_db() as conn:
        cursor = conn.cursor()
        
        # Modified archive_research table
        cursor.execute('''CREATE TABLE IF NOT EXISTS archive_research (
            id INT AUTO_INCREMENT PRIMARY KEY,
            archive_id INT NOT NULL,
            student_id VARCHAR(20) NOT NULL,
            department_id INT NOT NULL,
            course_id INT NOT NULL,
            project_title VARCHAR(255) NOT NULL,
            dateOfSubmit DATETIME NOT NULL,
            project_year VARCHAR(4) NOT NULL,
            project_abstract TEXT,
            keywords TEXT,
            content TEXT NOT NULL,
            research_owner_email VARCHAR(255),
            project_members TEXT,
            documents VARCHAR(255),
            file_size BIGINT UNSIGNED,
            page_count INT UNSIGNED,
            word_count INT UNSIGNED,
            character_count INT UNSIGNED
        )''')

        cursor.execute('''CREATE TABLE IF NOT EXISTS plagiarism_results (
            id INT AUTO_INCREMENT PRIMARY KEY,
            archive_id INT NOT NULL,
            submitted_sentence TEXT NOT NULL,
            existing_sentence TEXT NOT NULL,
            similar_archive_id INT NOT NULL,
            similarity_percentage DECIMAL(5,2) NOT NULL,
            is_plagiarized BOOLEAN NOT NULL,
            FOREIGN KEY (archive_id) REFERENCES archive_research(id),
            FOREIGN KEY (similar_archive_id) REFERENCES archive_research(id)
        )''')

        cursor.execute('''CREATE TABLE IF NOT EXISTS plagiarism_summary (
            id INT AUTO_INCREMENT PRIMARY KEY,
            archive_id INT NOT NULL,
            similar_archive_id INT NOT NULL,
            plagiarism_percentage DECIMAL(5,2) NOT NULL,
            FOREIGN KEY (archive_id) REFERENCES archive_research(id),
            FOREIGN KEY (similar_archive_id) REFERENCES archive_research(id)
        )''')
        
        conn.commit()

SUPPORTED_LANGUAGES = {
    'eng': 'English',
    'fil': 'Filipino/Tagalog',
    'ceb': 'Cebuano',
    'ilo': 'Ilocano',
    'hil': 'Hiligaynon',
    'bik': 'Bikol',
    'war': 'Waray',
    'pam': 'Kapampangan',
    'pan': 'Pangasinan'
}

def detect_language(text):
    filipino_markers = [
        'ng', 'mga', 'na', 'sa', 'ang', 'ay', 'at', 'po', 'ko', 'ka',
        'namin', 'natin', 'kami', 'tayo', 'siya', 'niya', 'nila', 'ninyo',
        'ako', 'bakit', 'ano', 'saan', 'paano', 'kailan', 'hindi', 'kasi'
    ]
    english_markers = [
        'the', 'is', 'are', 'and', 'to', 'it', 'of', 'in', 'on', 'for',
        'this', 'that', 'with', 'was', 'by', 'an', 'as', 'at', 'be'
    ]

    text_lower = text.lower()
    filipino_count = sum(1 for word in filipino_markers if word in text_lower.split())
    english_count = sum(1 for word in english_markers if word in text_lower.split())
    
    # Apply heuristic rules
    if filipino_count > 2 and filipino_count > english_count:
        return ['eng', 'fil']
    elif english_count > 2 and english_count > filipino_count:
        return ['eng']
    else:
        # If uncertain, return default languages
        return ['eng', 'fil']


def preprocess_image(image: Image.Image) -> Image.Image:
    """
    Preprocesses an image to enhance text readability for OCR.
    """
    try:
        os.makedirs('preprocessing_steps', exist_ok=True)

        if isinstance(image, Image.Image):
            image_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
        else:
            image_cv = image

        # Resize image
        target_height = 2000
        height, width = image_cv.shape[:2]
        scale_factor = target_height / height
        new_width = int(width * scale_factor)
        resized_image = cv2.resize(image_cv, (new_width, target_height), interpolation=cv2.INTER_CUBIC)
        cv2.imwrite('preprocessing_steps/1_resized.png', resized_image)

        # Convert to grayscale
        gray_image = cv2.cvtColor(resized_image, cv2.COLOR_BGR2GRAY)

        # Noise reduction
        denoised_image = cv2.fastNlMeansDenoising(gray_image, None, 30, 7, 21)
        cv2.imwrite('preprocessing_steps/2_denoised.png', denoised_image)

        # Apply CLAHE (Contrast Limited Adaptive Histogram Equalization)
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
        enhanced_image = clahe.apply(denoised_image)
        cv2.imwrite('preprocessing_steps/3_clahe.png', enhanced_image)

        # Adaptive thresholding (Tuned)
        adaptive_thresh = cv2.adaptiveThreshold(
            enhanced_image, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 15, 4
        )
        cv2.imwrite('preprocessing_steps/4_adaptive_thresh.png', adaptive_thresh)

        # Otsu’s Thresholding for better binarization
        _, otsu_thresh = cv2.threshold(enhanced_image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        cv2.imwrite('preprocessing_steps/5_otsu.png', otsu_thresh)

        # Morphological operations (Closing small gaps)
        kernel = np.ones((2, 2), np.uint8)
        morphed_image = cv2.morphologyEx(otsu_thresh, cv2.MORPH_CLOSE, kernel, iterations=2)
        cv2.imwrite('preprocessing_steps/6_morphology.png', morphed_image)

        # Sharpening filter for better edges
        sharpening_kernel = np.array([[-1, -1, -1], [-1, 9, -1], [-1, -1, -1]])
        sharpened_image = cv2.filter2D(morphed_image, -1, sharpening_kernel)
        cv2.imwrite('preprocessing_steps/7_sharpened.png', sharpened_image)

        # Invert colors for OCR
        final_image = cv2.bitwise_not(sharpened_image)
        cv2.imwrite('preprocessing_steps/8_final.png', final_image)

        return Image.fromarray(final_image)

    except Exception as e:
        print(f"Image preprocessing error: {e}")
        traceback.print_exc()
        return image
def correct_punctuation(text):
    """
    Corrects punctuation and spacing in the extracted text.
    """
    # Ensure spaces after punctuation
    text = re.sub(r'(?<=[.,])(?=[^\s])', r' ', text)
    
    # Ensure spaces before opening parentheses and after closing parentheses
    text = re.sub(r'(\()', r' \1', text)
    text = re.sub(r'(\))', r'\1 ', text)
    
    # Ensure spaces after commas
    text = re.sub(r'(?<=[^ ]),(?=[^ ])', r', ', text)
    
    # Ensure spaces after periods
    text = re.sub(r'(?<=[^ ])\.(?=[^ ])', r'. ', text)
    
    # Capitalize the first letter after a period
    text = re.sub(r'(?<=\. )([a-z])', lambda m: m.group(1).upper(), text)
    
    # Remove extra spaces
    text = re.sub(r'\s+', ' ', text)
    
    # Ensure proper spacing around hyphens
    text = re.sub(r'(?<=[^ ])-(?=[^ ])', r' - ', text)
    
    # Ensure proper spacing around slashes
    text = re.sub(r'(?<=[^ ])/(?=[^ ])', r' / ', text)
    
    # Remove spaces before punctuation
    text = re.sub(r'\s+([.,])', r'\1', text)
    
    if text and not text.endswith('.'):
        text += '.'
    
    return text.strip()

def extract_text_from_image(image, default_languages=['eng', 'fil']):
    """
    Enhanced text extraction with better handling of mixed languages and styled text.
    """
    try:
        processed_image = preprocess_image(image)
        
        custom_config = r'''--oem 3 --psm 6 --dpi 300 -c tessedit_char_blacklist=|[]{}()<> -c textord_heavy_nr=1 -c textord_min_linesize=2.5 -c textord_min_xheight=5 -c tessedit_enable_dict_correction=1 -c tessedit_enable_bigram_correction=1'''

        text_data = pytesseract.image_to_data(
            processed_image,
            lang='+'.join(default_languages),
            config=custom_config,
            output_type=pytesseract.Output.DICT
        )

        words = []
        confidences = []
        
        for i in range(len(text_data['text'])):
            word = text_data['text'][i].strip()
            conf = int(text_data['conf'][i])
            
            if word and conf > 60:
                words.append(word)
                confidences.append(conf)

        text = ' '.join(words)
        
        # Correct punctuation and spacing
        text = correct_punctuation(text)
        
        if text and not text.endswith('.') and text[-1].isalpha():
            text += '.'
        confidence = min(1.0, len(text) / 100) if text else 0.0
        return text, confidence

    except Exception as e:
        print(f"OCR Error: {str(e)}")
        return "", 0

def get_file_metrics(file_path, file_type):
    """Calculate file metrics including size, page count, word count, and character count"""
    metrics = {
        'file_size': os.path.getsize(file_path),
        'page_count': 0,
        'word_count': 0,
        'character_count': 0
    }
    
    content = ""
    
    # Default languages including Filipino
    default_languages = ['eng', 'fil']
    
    if file_type == 'application/pdf':
        pdf_document = fitz.open(file_path)
        metrics['page_count'] = len(pdf_document)
        all_texts = []
        all_scores = []
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Extract text directly from PDF
            page_text = page.get_text()
            content += page_text
            
            detected_langs = detect_language(page_text)
            
            # Only process images embedded in the PDF
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]
                    image = Image.open(io.BytesIO(image_bytes))
                    text, score = extract_text_from_image(image, detected_langs)
                    if text.strip():
                        content += f"\n{text}\n"
                except Exception as e:
                    print(f"Error processing PDF embedded image {img_index} on page {page_num}: {str(e)}")

        pdf_document.close()
            
    elif file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
        doc = Document(file_path)
        metrics['page_count'] = len(doc.paragraphs) // 40  # Approximate pages
        
        for para in doc.paragraphs:
            content += para.text + "\n"
        
        # Extract images from Word document
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    image = Image.open(io.BytesIO(image_data))
                    # Unpack the tuple returned by extract_text_from_image
                    image_text, score = extract_text_from_image(image, default_languages)
                    if image_text and image_text.strip():
                        content += f"\n{image_text}\n"
                except Exception as e:
                    print(f"Error processing image in Word document: {str(e)}")
    
    else:  # For text files
        with open(file_path, 'r') as file:
            content = file.read()
            metrics['page_count'] = len(content) // 3000  # Approximate pages
    
    # Calculate word and character counts
    metrics['word_count'] = len(content.split())
    metrics['character_count'] = len(content)
    
    return metrics, content

def normalize_text(text):
    # Convert to lowercase
    text = text.lower()

    # Remove leading and trailing spaces
    text = text.strip()
    
    # Replace unicode punctuation with standard punctuation
    text = text.translate(str.maketrans('', '', string.punctuation))
    
    # Remove multiple spaces between words
    text = re.sub(r'\s+', ' ', text)
    
    # Normalize hyphenated words (removes spaces around hyphens)    
    text = re.sub(r'\s*-\s*', '-', text)
    
    # Remove non-alphanumeric characters except spaces and hyphens
    text = re.sub(r'[^\w\s-]', '', text)

    # Handle redundant punctuation like multiple dots
    text = re.sub(r'\.\.+', '.', text)

    # Replace multiple consecutive hyphens with a single hyphen
    text = re.sub(r'-{2,}', '-', text)
    
    return text
    

def split_into_sentences(text):
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt')
    
    # First, preserve certain abbreviations and numbers
    preserved_text = text
    preserved_text = re.sub(r'(?<=\d)\.(?=\d)', '@DOT@', preserved_text)  # Preserve decimal points
    preserved_text = re.sub(r'(?<=\w)\.(?=\w)', '@DOT@', preserved_text)  # Preserve abbreviations
    
    # Split into sentences
    sentences = sent_tokenize(preserved_text)
    
    # Restore preserved characters
    sentences = [s.replace('@DOT@', '.') for s in sentences]
    
    # Additional cleaning and filtering
    cleaned_sentences = []
    for sentence in sentences:
        # Normalize the sentence
        sentence = normalize_text(sentence)
        
        # Count actual words (excluding numbers and single characters)
        word_count = len([w for w in sentence.split() if re.match(r'[a-zA-Z]{2,}', w)])
        
        # Filter sentences based on word count and other criteria
        if (7 <= word_count <= 60 and  # Word count between 7 and 50
            len(sentence) >= 40 and     # Minimum character length
            not re.match(r'^[^a-zA-Z]*$', sentence) and  # Contains letters
            not re.match(r'^(table|figure|fig)', sentence.lower())):  # Not a table/figure reference
            cleaned_sentences.append(sentence)
    
    return cleaned_sentences

def calculate_similarity(submitted_content, submissions, similarity_threshold=0.7):
    submitted_sentences = [normalize_text(sentence) for sentence in split_into_sentences(submitted_content)]

    # Prepare a list of sentences and their corresponding submission IDs
    all_existing_sentences = []
    all_existing_submission_ids = []

    for submission_id, content in submissions:
        sentences = [normalize_text(sentence) for sentence in split_into_sentences(content)]
        all_existing_sentences.extend(sentences)
        all_existing_submission_ids.extend([submission_id] * len(sentences))  # Match each sentence with its submission ID

    # Combine submitted and existing sentences
    all_sentences = submitted_sentences + all_existing_sentences

    if not all_sentences:
        return []

    # Vectorize the sentences
    vectorizer = TfidfVectorizer(lowercase=True, stop_words=None, token_pattern=r'\b\w+\b').fit_transform(all_sentences)

    # Similarity between submitted sentences and all existing ones
    similarity_matrix = cosine_similarity(vectorizer[:len(submitted_sentences)], vectorizer[len(submitted_sentences):])

    # Find matching sentences
    matching_results = []
    for i, submitted_sentence in enumerate(submitted_sentences):
        for j, existing_sentence in enumerate(all_existing_sentences):
            similarity_score = similarity_matrix[i][j]
            if similarity_score > similarity_threshold:
                matching_results.append({
                    'submitted_sentence': submitted_sentence,
                    'existing_sentence': existing_sentence,
                    'similarity_percentage': similarity_score * 100,
                    'similar_archive_id': all_existing_submission_ids[j]
                })

    return matching_results

@app.route('/')
def index():    
    return send_from_directory('.', 'student/project_list.php')

@app.route('/upload_research', methods=['POST'])
def upload_file():
    print("Form Data:", request.form)
    print("Files Data:", request.files)
    if 'file' not in request.files:
        return jsonify({'status': 'error', 'message': "No file part"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'status': 'error', 'message': "No selected file"}), 400

    # Extract form data
    archive_id = request.form.get('archive_id')
    student_id = request.form.get('student_id')
    project_title = request.form.get('project_title')
    project_year = request.form.get('year')
    department_id = request.form.get('department_id')
    course_id = request.form.get('course_id')
    abstract = request.form.get('abstract')
    keywords = request.form.get('keywords')
    project_members = request.form.get('project_members')
    pdf_path = request.form.get('pdf_path')
    owner_email = request.form.get('owner_email')

    if not student_id:
        return jsonify({'status': 'error', 'message': "No student ID provided"}), 400

    try:
        # Step 1: Save the file
        upload_folder = os.path.join(os.getcwd(), 'pdf_files')
        os.makedirs(upload_folder, exist_ok=True)
        file_path = os.path.join(upload_folder, file.filename)
        # File metrics and content
        file_metrics, content = get_file_metrics(file_path, file.content_type)
        if not content:
            return jsonify({'status': 'error', 'message': "Failed to extract content from file"}), 400
        print(content)
        # Save the uploaded file details in the database
        manila_tz = pytz.timezone('Asia/Manila')
        current_time = datetime.now(manila_tz)
        
        with get_db() as conn:
            cursor = conn.cursor()
            # First, insert the research with initial status
            cursor.execute("""
                INSERT INTO archive_research (
                    archive_id, student_id, department_id, course_id,
                    project_title, dateOfSubmit, project_year, project_abstract,
                    keywords, content, research_owner_email, project_members,
                    documents, file_size, page_count, word_count, character_count,
                    document_status  -- Added initial status
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    archive_id, student_id, department_id, course_id,
                    project_title, current_time, project_year, abstract,
                    keywords, content, owner_email, project_members,
                    pdf_path, file_metrics['file_size'], file_metrics['page_count'],
                    file_metrics['word_count'], file_metrics['character_count'],
                    'Processing'  # Initial status while plagiarism check runs
                ))
            conn.commit()
            new_archive_id = cursor.lastrowid

            # Get department info for response
            cursor.execute("SELECT * FROM departments WHERE id = %s", (department_id,))
            departments_db = cursor.fetchall()
            department = departments_db[0][2]

        # Prepare response before starting plagiarism check
        response = {
            'status': 'success',
            'department': department,
            'document_status': 'Processing',
            'file_metrics': file_metrics,
            'archive_id': archive_id,
            'message': 'File uploaded successfully. Plagiarism check is running in the background.'
        }

        plagiarism_check(new_archive_id, student_id, owner_email, content, current_time)

        return jsonify(response)

    except Exception as e:
        print(f"Error: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

def plagiarism_check(new_archive_id, student_id, owner_email, content, current_time):
    """Separate function for background plagiarism checking"""
    try:
        with get_db() as conn:
            cursor = conn.cursor()
        
            cursor.execute("""
                UPDATE archive_research 
                SET document_status = 'Checking Plagiarism' 
                WHERE id = %s
            """, (new_archive_id,))
            conn.commit()

            # Fetch existing submissions for comparison
            cursor.execute("SELECT id, content FROM archive_research WHERE student_id != %s AND (research_owner_email = %s OR project_members NOT LIKE %s)", (student_id, owner_email, f"%{owner_email}%"))
            submissions = cursor.fetchall()

            plagiarism_results = {}
            if submissions:
                matching_sentences = calculate_similarity(content, submissions)
                
                # Track unique matched sentences per document
                doc_matches = {}
                
                for match in matching_sentences:
                    similar_id = match['similar_archive_id']
                    
                    if similar_id not in doc_matches:
                        doc_matches[similar_id] = set()
                    
                    # Add the matched sentence to track uniqueness
                    match_key = (match['submitted_sentence'], match['existing_sentence'])
                    doc_matches[similar_id].add(match_key)
                    
                    cursor.execute("""
                        INSERT INTO plagiarism_results (
                            archive_id, similar_archive_id, submitted_sentence,
                            existing_sentence, similarity_percentage, is_plagiarized
                        ) VALUES (%s, %s, %s, %s, %s, %s)
                        """, (
                        new_archive_id, similar_id, match['submitted_sentence'],
                        match['existing_sentence'], match['similarity_percentage'], True
                    ))
                    
                    if similar_id not in plagiarism_results:
                        cursor.execute("SELECT content FROM archive_research WHERE id = %s", (similar_id,))
                        similar_content = cursor.fetchone()[0]
                        total_sentences = len(split_into_sentences(similar_content))
                        
                        plagiarism_results[similar_id] = {
                            'matched_sentences': 0,
                            'total_similarity': 0,
                            'total_sentences': total_sentences
                        }
                    
                    # Only increment matched_sentences once per unique match
                    if len(doc_matches[similar_id]) > plagiarism_results[similar_id]['matched_sentences']:
                        plagiarism_results[similar_id]['matched_sentences'] = len(doc_matches[similar_id])
                    plagiarism_results[similar_id]['total_similarity'] += match['similarity_percentage']

                print(plagiarism_results)

                for similar_id, results in plagiarism_results.items():
                    avg_similarity = results['total_similarity'] / results['matched_sentences']
                    plagiarism_percentage = (results['total_similarity'] / results['total_sentences']) if results['total_sentences'] > 0 else 0

                    cursor.execute("""
                        INSERT INTO plagiarism_summary (
                            archive_id, similar_archive_id, plagiarism_percentage
                        ) VALUES (%s, %s, %s)
                        """, (new_archive_id, similar_id, plagiarism_percentage))
            conn.commit()

            print(plagiarism_results)
            # Calculate final plagiarism results and update status
            cursor.execute("""
                SELECT COUNT(plagiarism_summary.id) as total_ids, 
                       SUM(plagiarism_percentage) as total_percentage 
                FROM plagiarism_summary 
                WHERE archive_id = %s
                GROUP BY archive_id
                """, (new_archive_id,))
            
            result = cursor.fetchone()
            sum_of_percentage = result[1] if result else 0
            document_status = "Accepted" if sum_of_percentage <= 20 else "Rejected"
            date_publish = current_time.strftime('%Y-%m-%d') if document_status == "Accepted" else None

            cursor.execute("""
                UPDATE archive_research 
                SET document_status = %s, 
                    date_published = %s 
                WHERE id = %s
                """, (document_status, date_publish, new_archive_id))
            conn.commit()

    except Exception as e:
        print(f"Error during plagiarism check: {str(e)}")

@app.route('/check_db_connection', methods=['GET'])
def check_db_connection():
    try:
        with get_db() as conn:
            conn.execute("SELECT 1")
        return jsonify(status="Database connection successful"), 200
    except Exception as e:
        return jsonify(status="Database connection failed", error=str(e)), 500
    
if __name__ == '__main__':
    init_db()
    app.run(debug=True, port=3000)